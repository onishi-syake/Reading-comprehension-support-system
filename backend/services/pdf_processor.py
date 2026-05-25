import pdfplumber
import pytesseract
from pdf2image import convert_from_bytes
from docx import Document
from fastapi import UploadFile
import io
import unicodedata

class DocumentProcessor:
    """PDF・Wordファイル処理クラス"""

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt'}

    async def extract_text(self, file: UploadFile) -> dict:
        """
        アップロードされたファイルからテキストを抽出

        Args:
            file: アップロードされたPDF or Wordファイル

        Returns:
            {"text": 抽出テキスト, "method": 抽出方法}
        """
        filename = file.filename or ""
        ext = self._get_extension(filename)

        if ext == '.txt':
            return await self._extract_from_txt(file)
        elif ext == '.docx':
            return await self._extract_from_docx(file)
        elif ext == '.pdf':
            return await self._extract_from_pdf(file)
        else:
            raise Exception(f"未対応のファイル形式です: {ext}")

    def _get_extension(self, filename: str) -> str:
        """ファイル名から拡張子を取得"""
        import os
        return os.path.splitext(filename)[1].lower()

    async def _extract_from_pdf(self, file: UploadFile) -> dict:
        """PDFからテキストを抽出"""
        try:
            content = await file.read()
            pdf_file = io.BytesIO(content)

            extracted_text = ""
            with pdfplumber.open(pdf_file) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        extracted_text += text

            if not self._needs_ocr(extracted_text, page_count):
                return {"text": self.clean_text(extracted_text), "method": "pdfplumber"}

            ocr_text = self._extract_with_ocr(content)
            return {"text": self.clean_text(ocr_text), "method": "ocr"}

        except Exception as e:
            raise Exception(f"PDFテキスト抽出エラー: {str(e)}")

    async def _extract_from_docx(self, file: UploadFile) -> dict:
        """Wordファイルからテキストを抽出（段落・テーブル両方）"""
        try:
            content = await file.read()
            doc = Document(io.BytesIO(content))

            texts = []

            # 段落からテキスト取得
            for p in doc.paragraphs:
                if p.text.strip():
                    texts.append(p.text)

            # テーブルからテキスト取得（重複セルを除外）
            for table in doc.tables:
                for row in table.rows:
                    seen = set()
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and cell_text not in seen:
                            seen.add(cell_text)
                            texts.append(cell_text)

            extracted_text = "\n\n".join(texts)

            return {"text": self.clean_text(extracted_text), "method": "docx"}

        except Exception as e:
            raise Exception(f"Wordテキスト抽出エラー: {str(e)}")

    async def _extract_from_txt(self, file: UploadFile) -> dict:
        """テキストファイルからテキストを抽出"""
        try:
            content = await file.read()
            extracted_text = content.decode('utf-8')
            return {"text": self.clean_text(extracted_text), "method": "txt"}
        except UnicodeDecodeError:
            # UTF-8で読めない場合はShift_JISを試行
            extracted_text = content.decode('shift_jis')
            return {"text": self.clean_text(extracted_text), "method": "txt"}
        except Exception as e:
            raise Exception(f"テキストファイル読み込みエラー: {str(e)}")

    def _needs_ocr(self, text: str, page_count: int) -> bool:
        """OCRが必要か判定"""
        if not text.strip():
            return True

        if len(text.strip()) / max(page_count, 1) < 100:
            return True

        garbled_count = sum(
            1 for c in text
            if c == '\ufffd'
            or (unicodedata.category(c).startswith('C') and c not in '\n\r\t')
        )
        if len(text) > 0 and garbled_count / len(text) > 0.1:
            return True

        return False

    def _extract_with_ocr(self, pdf_bytes: bytes) -> str:
        """OCRでテキスト抽出"""
        images = convert_from_bytes(pdf_bytes, dpi=300)
        extracted_text = ""
        for image in images:
            text = pytesseract.image_to_string(image, lang='eng+jpn')
            if text:
                extracted_text += text
        return extracted_text

    def clean_text(self, text: str) -> str:
        """
        抽出したテキストをクリーニング

        Args:
            text: 元のテキスト

        Returns:
            クリーニングされたテキスト
        """
        import re

        # 改行の正規化
        text = text.replace('\r\n', '\n')

        # 単一の改行を削除（段落途中の改行を結合）
        text = re.sub(r'(?<!\n)\n(?!\n)', '', text)

        # 複数の改行を1つに
        text = re.sub(r'\n{2,}', '\n', text)

        # 複数の空白を1つに
        text = re.sub(r' +', ' ', text)

        return text.strip()

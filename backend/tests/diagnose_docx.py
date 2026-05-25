"""
Word文書の構造を診断するスクリプト
テキストがどの要素に格納されているか調査する

使い方:
    python tests/diagnose_docx.py <Wordファイルパス>
"""

from docx import Document
import sys
import os


def main():
    if len(sys.argv) < 2:
        print("使い方: python tests/diagnose_docx.py <Wordファイルパス>")
        sys.exit(1)

    docx_path = sys.argv[1]
    doc = Document(docx_path)

    # 1. 段落
    print("=" * 60)
    print(f"【1. 段落 (paragraphs)】: {len(doc.paragraphs)}個")
    print("=" * 60)
    for i, p in enumerate(doc.paragraphs[:10]):
        if p.text.strip():
            print(f"  [{i}] {p.text[:80]}")

    # 2. 表（テーブル）
    print(f"\n{'=' * 60}")
    print(f"【2. テーブル (tables)】: {len(doc.tables)}個")
    print("=" * 60)
    for t_idx, table in enumerate(doc.tables):
        print(f"\n  テーブル[{t_idx}]: {len(table.rows)}行 x {len(table.columns)}列")
        for r_idx, row in enumerate(table.rows[:5]):
            cells = [cell.text[:30] for cell in row.cells]
            print(f"    行[{r_idx}]: {cells}")
        if len(table.rows) > 5:
            print(f"    ... 残り{len(table.rows) - 5}行")

    # 3. ヘッダー・フッター
    print(f"\n{'=' * 60}")
    print("【3. ヘッダー・フッター】")
    print("=" * 60)
    for i, section in enumerate(doc.sections):
        header_text = section.header.paragraphs[0].text if section.header.paragraphs else ""
        footer_text = section.footer.paragraphs[0].text if section.footer.paragraphs else ""
        print(f"  セクション[{i}] ヘッダー: {header_text[:50]}")
        print(f"  セクション[{i}] フッター: {footer_text[:50]}")

    # 4. 全テキスト（XML直接解析）
    print(f"\n{'=' * 60}")
    print("【4. XML内の全テキスト（先頭500文字）】")
    print("=" * 60)
    from docx.oxml.ns import qn
    all_text = ""
    for elem in doc.element.iter(qn('w:t')):
        all_text += elem.text or ""
    print(f"  総文字数: {len(all_text)}")
    print(f"  先頭500文字: {all_text[:500]}")


if __name__ == "__main__":
    main()
